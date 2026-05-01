#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "python-docx>=1.1.0",
#     "markdown>=3.5",
# ]
# ///
"""docx <-> markdown conversion tool.

Usage:
    uv run scripts/docx_convert.py input.docx                # docx -> md
    uv run scripts/docx_convert.py input.docx -o output.md   # docx -> md (custom output)
    uv run scripts/docx_convert.py input.md                  # md -> docx
    uv run scripts/docx_convert.py input.md -o output.docx   # md -> docx (custom output)
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def docx_to_md(docx_path: Path) -> str:
    """Convert a .docx file to markdown text."""
    doc = Document(str(docx_path))
    lines: list[str] = []

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            if not para.text.strip():
                lines.append("")
                continue

            style = para.style.name.lower() if para.style else ""

            if style.startswith("heading"):
                level = _heading_level(style)
                prefix = "#" * level + " "
                content = _para_to_md(para)
                lines.append(f"{prefix}{content}")
            elif style.startswith("list"):
                content = _para_to_md(para)
                if "number" in style:
                    lines.append(f"1. {content}")
                else:
                    lines.append(f"- {content}")
            else:
                content = _para_to_md(para)
                lines.append(content)

        elif tag == "tbl":
            table_idx = sum(
                1 for c in body[:body.index(child) + 1]
                if c.tag.split("}")[-1] == "tbl"
            )
            if table_idx <= len(doc.tables):
                table = doc.tables[table_idx - 1]
                lines.append("")
                lines.append(_table_to_md(table))
                lines.append("")

    while len(lines) > 1 and lines[-1] == "":
        lines.pop()

    result = "\n".join(lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result


def _para_to_md(para) -> str:
    """Convert a paragraph's runs to markdown text."""
    if not para.runs:
        return para.text

    parts: list[str] = []
    for run in para.runs:
        if not run.text:
            continue

        text = run.text
        is_bold = run.bold
        is_italic = run.italic
        is_strike = getattr(run, "strike", None)
        if is_strike is None:
            rPr = run._element.rPr if hasattr(run, "_element") else None
            if rPr is not None:
                strike_elem = rPr.find(qn("w:strike"))
                is_strike = strike_elem is not None and strike_elem.get(qn("w:val")) not in ("0", "false")
            else:
                is_strike = False

        if is_strike:
            text = f"~~{text}~~"
        if is_bold:
            text = f"**{text}**"
        if is_italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts) if parts else para.text


def md_to_docx(md_path: Path, out_path: Path) -> None:
    """Convert a markdown file to .docx."""
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_separator(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            if table_rows:
                _write_table(doc, table_rows)
            continue

        if not line:
            i += 1
            continue

        if line.startswith("#"):
            level = _count_heading(line)
            content = line.lstrip("#").strip()
            p = doc.add_heading(level=level)
            _add_inline_md_to_paragraph(p, content)
        elif line.startswith("- ") or line.startswith("* "):
            content = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_md_to_paragraph(p, content)
        elif re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline_md_to_paragraph(p, content)
        elif line.startswith("> "):
            content = line[2:]
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(content)
            run.italic = True
        else:
            p = doc.add_paragraph()
            _add_inline_md_to_paragraph(p, line)

        i += 1

    doc.save(str(out_path))


def _add_inline_md_to_paragraph(paragraph, text: str) -> None:
    """Parse markdown inline formatting and add to paragraph."""
    pattern = r"(\*\*.*?\*\*|~~.*?~~|`.*?`|\*.*?\*|\[.*?\]\(.*?\))"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("~~") and part.endswith("~~"):
            run = paragraph.add_run(part[2:-2])
            run.strike = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            bracket_end = part.index("](")
            link_text = part[1:bracket_end]
            url = part[bracket_end + 2:-1]
            run = paragraph.add_run(link_text)
            run.font.color.rgb = None
            rPr = run._element.get_or_add_rPr()
            hlink = rPr.makeelement(qn("w:highlight"), {"w:val": "blue"})
            rPr.append(hlink)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _heading_level(style_name: str) -> int:
    m = re.search(r"(\d)", style_name)
    return int(m.group(1)) if m else 1


def _table_to_md(table) -> str:
    rows = table.rows
    if not rows:
        return ""

    header = [cell.text for cell in rows[0].cells]
    sep = ["---"] * len(header)
    data = [[cell.text for cell in row.cells] for row in rows[1:]]

    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(sep) + " |")
    for row in data:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def _count_heading(line: str) -> int:
    m = re.match(r"^(#+)", line)
    return len(m.group(1)) if m else 1


def _write_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            t.cell(i, j).text = cell_text


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert between docx and markdown formats."
    )
    parser.add_argument("input", type=Path, help="Input file (.docx or .md)")
    parser.add_argument(
        "-o", "--output", type=Path, default=None, help="Output file path"
    )
    args = parser.parse_args()

    inp: Path = args.input
    if not inp.exists():
        print(f"Error: file not found: {inp}", file=sys.stderr)
        sys.exit(1)

    suffix = inp.suffix.lower()

    if suffix == ".docx":
        out = args.output or inp.with_suffix(".md")
        md_text = docx_to_md(inp)
        out.write_text(md_text, encoding="utf-8")
        print(f"Converted: {inp} -> {out}")
    elif suffix in (".md", ".markdown"):
        out = args.output or inp.with_suffix(".docx")
        md_to_docx(inp, out)
        print(f"Converted: {inp} -> {out}")
    else:
        print(f"Error: unsupported file type '{suffix}'. Use .docx or .md", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
